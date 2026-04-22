from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "OFFICE_HANDOFF_PRINT.docx"


def set_font(run_or_style, name: str, size: int | None = None, bold: bool | None = None):
    font = run_or_style.font
    font.name = name
    if hasattr(run_or_style, "_element"):
        run_or_style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold


def add_table(document: Document, headers: list[str], rows: list[tuple[str, str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    set_font(normal, "Microsoft JhengHei", size=11)

    heading1 = doc.styles["Heading 1"]
    set_font(heading1, "Microsoft JhengHei", size=18, bold=True)

    heading2 = doc.styles["Heading 2"]
    set_font(heading2, "Microsoft JhengHei", size=13, bold=True)

    heading3 = doc.styles["Heading 3"]
    set_font(heading3, "Microsoft JhengHei", size=11, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("飯店合作夥伴入場管理系統\n辦公室操作 SOP")
    set_font(title_run, "Microsoft JhengHei", size=20, bold=True)

    doc.add_paragraph("適用情境：固定單機 Demo 版操作、辦公室交接、關機後重開。")

    doc.add_heading("用途說明", level=2)
    for item in [
        "這是一套固定在本機使用的 Demo 版「飯店合作夥伴入場管理系統」。",
        "系統固定開在這一台電腦。",
        "資料保存在這台電腦的瀏覽器 localStorage。",
        "同事直接使用這台電腦操作。",
        "目前唯一正式使用網址為 http://localhost:5179/ 。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("每日啟動 SOP", level=2)
    add_table(
        doc,
        ["步驟", "操作"],
        [
            ("1", "打開 PowerShell"),
            ("2", r"輸入 cd D:\hotel-partner-entry-system\web"),
            ("3", "輸入 npm run dev"),
            ("4", "等終端機顯示網址後，用瀏覽器打開"),
            ("5", "一律使用 http://localhost:5179/"),
        ],
    )

    doc.add_heading("如果網站已經開著", level=2)
    for step in [
        "直接打開瀏覽器。",
        "輸入 http://localhost:5179/ 。",
        "開始操作。",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("不幸關機後的重新開啟方式", level=2)

    doc.add_heading("情境 A：電腦重開機了，網站打不開", level=3)
    for step in [
        "打開 PowerShell。",
        r"輸入 cd D:\hotel-partner-entry-system\web 。",
        "再輸入 npm run dev 。",
        "用瀏覽器打開 http://localhost:5179/ 。",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("情境 B：網站頁面開著，但突然無法連線", level=3)
    for step in [
        "PowerShell 視窗是不是被關掉了。",
        "電腦是不是剛重開機。",
        "現在開的是不是 http://localhost:5179/ 。",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph("如果不確定，最簡單做法：")
    for step in [
        "關掉舊分頁。",
        "重新打開 PowerShell。",
        r"重新輸入 cd D:\hotel-partner-entry-system\web 。",
        "再輸入 npm run dev 。",
        "改用新的 http://localhost:5179/ 網址。",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("正確使用規則", level=2)
    add_table(
        doc,
        ["情境", "正確做法"],
        [
            ("要保留資料", "只用 5179"),
            ("同事要操作", "直接用這台電腦"),
            ("要保留資料備份", "依照 EXPORT_DATA.md 匯出 JSON"),
            ("QR 掃不出來", "改用序號手動搜尋"),
        ],
    )

    doc.add_heading("Demo 操作順序", level=2)
    for step in [
        "到「發放憑證」。",
        "建立新的 QR。",
        "下載 QR 圖檔。",
        "到「現場核銷」。",
        "掃碼或用序號手動查找。",
        "輸入本次核銷人數。",
        "到「管理後台」查看統計變化。",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("常見問題", level=2)
    doc.add_heading("為什麼資料會保留？", level=3)
    doc.add_paragraph("因為資料存在這台電腦的瀏覽器裡，只要不清除瀏覽器資料，就會保留。")

    doc.add_heading("如果同事要接手怎麼做？", level=3)
    doc.add_paragraph("直接來這台電腦操作，不需要另外安裝或複製專案。")

    doc.add_heading("如果 QR 掃不出來怎麼辦？", level=3)
    doc.add_paragraph("請先改用 QR 圖上的序號，在「現場核銷」頁手動搜尋對應憑證。")

    doc.add_heading("重要提醒", level=2)
    for item in [
        "正式操作請只用 5179。",
        "不要清除瀏覽器資料。",
        "不要換瀏覽器。",
        "關機後網站不會自動重開，需要重新執行 npm run dev。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("相關文件", level=2)
    for item in [
        "web/DEMO_SCRIPT.md",
        "web/EXPORT_DATA.md",
        "web/OFFICE_HANDOFF.md",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
